from docx import Document
import os

def replace_partial_labels(doc_path, output_path, original_text, new_labels):
    doc = Document(doc_path)

    label_index = 0  # Track how many labels we’ve replaced

    # Replace in all paragraphs
    for para in doc.paragraphs:
        for run in para.runs:
            if original_text in run.text:
                if label_index < len(new_labels):
                    run.text = run.text.replace(original_text, new_labels[label_index])
                    label_index += 1
                else:
                    run.text = run.text.replace(original_text, "")  # blank out

    # Replace in tables too (if needed)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if original_text in run.text:
                            if label_index < len(new_labels):
                                run.text = run.text.replace(original_text, new_labels[label_index])
                                label_index += 1
                            else:
                                run.text = run.text.replace(original_text, "")

    doc.save(output_path)
    print(f"✅ Created: {output_path} with {label_index} replaced labels.")

# File paths
template_path = r"D:\Code\LABEL CHANGER\STICKER.docx"
output_path = r"D:\Code\LABEL CHANGER\newsticker.docx"

# The label text you want to replace
original_text = "LABEL"

# The new labels (only these many will be filled; rest go blank)
new_labels = []

while True:
    label = input("\nEnter new label (or leave blank to finish): ").strip()
    if not label:
        break
    try:
        qty = int(input(f"How many times to repeat '{label}'? "))
        new_labels.extend([label] * qty)
    except ValueError:
        print("❌ Please enter a valid number.")

# Confirm and replace
print(f"\n🔁 Total labels to insert: {len(new_labels)}")
replace_partial_labels(template_path, output_path, original_text, new_labels)